#!/usr/bin/env python3
"""
创建高三数学复习规划 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_el = tcPr.find(qn(f'w:{edge}'))
            if edge_el is None:
                edge_el = docx.oxml.OxmlElement(f'w:{edge}')
                tcPr.append(edge_el)
            edge_el.set(qn('w:val'), kwargs[edge])

def create_math_review_plan():
    doc = Document()
    
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading('高三数学复习规划', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('—— 科学备考，高效提分')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # 一、复习目标
    doc.add_heading('一、复习目标', 1)
    goals = [
        '夯实基础：全面掌握高中数学核心知识点，构建完整的知识体系',
        '提升能力：培养数学思维能力，提高解题速度和准确率',
        '突破难点：攻克函数、解析几何、导数等重难点专题',
        '应试技巧：熟悉高考题型，掌握答题策略和时间分配'
    ]
    for goal in goals:
        p = doc.add_paragraph(goal, style='List Bullet')
        p.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # 二、复习阶段规划
    doc.add_heading('二、复习阶段规划', 1)
    
    # 阶段表格
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    headers = ['阶段', '时间', '主要内容', '目标']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
    
    # 第一阶段：一轮复习
    row = table.rows[1]
    row.cells[0].text = '一轮复习\n（基础夯实）'
    row.cells[1].text = '高三上学期\n9月-12月'
    row.cells[2].text = '• 集合与常用逻辑用语\n• 函数与导数\n• 三角函数与解三角形\n• 数列\n• 立体几何\n• 解析几何'
    row.cells[3].text = '全面梳理知识点，\n建立知识框架，\n基础题不丢分'
    
    # 第二阶段：二轮复习
    row = table.rows[2]
    row.cells[0].text = '二轮复习\n（专题突破）'
    row.cells[1].text = '高三下学期\n1月-3月'
    row.cells[2].text = '• 函数与方程思想\n• 数形结合思想\n• 分类讨论思想\n• 转化与化归思想\n• 选填题专项\n• 解答题专项'
    row.cells[3].text = '突破重难点，\n掌握解题方法，\n提高综合能力'
    
    # 第三阶段：三轮复习
    row = table.rows[3]
    row.cells[0].text = '三轮复习\n（模拟冲刺）'
    row.cells[1].text = '4月-5月'
    row.cells[2].text = '• 全真模拟考试\n• 错题回顾\n• 应试技巧训练\n• 心理调适'
    row.cells[3].text = '适应考试节奏，\n查漏补缺，\n稳定心态'
    
    # 第四阶段：考前调整
    row = table.rows[4]
    row.cells[0].text = '考前调整\n（回归基础）'
    row.cells[1].text = '6月初'
    row.cells[2].text = '• 回归课本\n• 回顾错题\n• 保持手感\n• 调整作息'
    row.cells[3].text = '巩固基础，\n保持状态，\n自信应考'
    
    doc.add_paragraph()
    
    # 三、各模块复习要点
    doc.add_heading('三、各模块复习要点', 1)
    
    modules = [
        ('1. 函数与导数（约30分）', [
            '掌握函数的定义域、值域、单调性、奇偶性、周期性',
            '熟练运用导数研究函数的单调性、极值、最值',
            '重点突破：函数零点问题、不等式证明、参数讨论'
        ]),
        ('2. 三角函数与解三角形（约17分）', [
            '熟记三角函数的图像与性质、诱导公式',
            '掌握正弦定理、余弦定理的应用',
            '重点突破：三角恒等变换、解三角形的实际应用'
        ]),
        ('3. 数列（约10分）', [
            '掌握等差数列、等比数列的通项公式与求和公式',
            '理解数列的递推关系',
            '重点突破：数列求和方法（裂项相消、错位相减）'
        ]),
        ('4. 立体几何（约22分）', [
            '掌握空间几何体的结构特征、三视图',
            '熟练运用向量法解决空间角问题',
            '重点突破：空间平行与垂直的证明'
        ]),
        ('5. 解析几何（约22分）', [
            '掌握直线、圆、椭圆、双曲线、抛物线的方程与性质',
            '熟练运用韦达定理处理弦长、中点问题',
            '重点突破：圆锥曲线的综合应用'
        ]),
        ('6. 概率与统计（约17分）', [
            '掌握古典概型、几何概型',
            '理解离散型随机变量的分布列、期望与方差',
            '重点突破：回归分析、独立性检验'
        ])
    ]
    
    for title, points in modules:
        doc.add_heading(title, 2)
        for point in points:
            p = doc.add_paragraph(point, style='List Bullet')
            p.runs[0].font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 四、每日学习计划
    doc.add_heading('四、每日学习计划（建议）', 1)
    
    schedule_table = doc.add_table(rows=5, cols=2)
    schedule_table.style = 'Light List Accent 1'
    
    schedule_data = [
        ('时间段', '学习内容'),
        ('早晨（30分钟）', '背诵公式、定理、概念\n回顾昨日错题'),
        ('上午课堂', '认真听讲，做好笔记\n紧跟老师思路'),
        ('下午自习（1小时）', '完成当日作业\n针对性练习薄弱专题'),
        ('晚上（1-1.5小时）', '整理错题本\n刷高考真题/模拟题\n总结解题方法')
    ]
    
    for i, (time, content) in enumerate(schedule_data):
        row = schedule_table.rows[i]
        row.cells[0].text = time
        row.cells[1].text = content
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # 五、复习策略
    doc.add_heading('五、复习策略与技巧', 1)
    
    strategies = [
        ('1. 错题本策略', [
            '建立分类错题本（按知识点、错误类型分类）',
            '定期回顾（每周、每月）',
            '分析错误原因，总结解题思路',
            '对反复出错的题目进行专项突破'
        ]),
        ('2. 真题训练策略', [
            '近5年高考真题至少做3遍',
            '第一遍：限时模拟，检验水平',
            '第二遍：分类研究，总结规律',
            '第三遍：查漏补缺，巩固提升'
        ]),
        ('3. 时间分配策略', [
            '选填题（40-50分钟）：先易后难，遇到难题标记后跳过',
            '解答题（70-80分钟）：确保前3题满分，后3题争取步骤分',
            '检查（10-15分钟）：重点检查选填题和计算步骤'
        ]),
        ('4. 心态调整策略', [
            '保持规律作息，保证充足睡眠',
            '适当运动，缓解压力',
            '积极自我暗示，建立自信',
            '与同学、老师多交流，及时解决困惑'
        ])
    ]
    
    for title, points in strategies:
        doc.add_heading(title, 2)
        for point in points:
            p = doc.add_paragraph(point, style='List Bullet')
            p.runs[0].font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 六、推荐资料
    doc.add_heading('六、推荐复习资料', 1)
    
    resources = [
        '教材：人教版高中数学必修+选修（回归课本，夯实基础）',
        '真题：《五年高考三年模拟》《高考数学真题全刷》',
        '专题：《高考数学压轴题破解策略》《导数与圆锥曲线专题》',
        '工具：错题本、公式手册、思维导图'
    ]
    
    for resource in resources:
        p = doc.add_paragraph(resource, style='List Bullet')
        p.runs[0].font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 七、寄语
    doc.add_heading('七、备考寄语', 1)
    
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_run = quote.add_run(
        '数学是一门需要积累和思考的学科，\n'
        '没有捷径，但有方法。\n'
        '坚持每日练习，善于总结反思，\n'
        '相信通过科学规划和不懈努力，\n'
        '你一定能在高考中取得理想成绩！\n\n'
        '加油！💪'
    )
    quote_run.font.size = Pt(12)
    quote_run.font.color.rgb = RGBColor(0, 102, 204)
    quote_run.font.italic = True
    
    # 保存文档
    output_path = '/root/.openclaw/workspace/高三数学复习规划.docx'
    doc.save(output_path)
    print(f'✅ Word 文档已生成！')
    print(f'📁 文件位置: {output_path}')
    return output_path

if __name__ == '__main__':
    create_math_review_plan()
